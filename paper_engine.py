"""
Theophysics Paper Engine
========================
Convert markdown papers to styled Word documents and PDFs.

Usage:
  python paper_engine.py                    # Convert all FT-*.md / FP-*.md
  python paper_engine.py paper.md           # Convert specific file
  python paper_engine.py --pdf              # Also export to PDF
  python paper_engine.py --templates        # Generate 7Q score templates
  python paper_engine.py --dir /path/to/md  # Process from specific directory

Author: David Lowe | Theophysics Research
"""
import os
import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════
# COLOR PALETTE
# ═══════════════════════════════════════════════════════════════
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_GRAY = RGBColor(0x55, 0x55, 0x55)
MED_GRAY = RGBColor(0x88, 0x88, 0x88)
BODY_BLACK = RGBColor(0x00, 0x00, 0x00)
HEADING_COLOR = RGBColor(0x1A, 0x3C, 0x5E)
ACCENT_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
BLOCKQUOTE_GRAY = RGBColor(0x44, 0x44, 0x44)

MARGIN = Inches(1)


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def set_run(run, font_name, size, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, font_name, size, bold=False, italic=False,
                  color=None, alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, font_name, size, bold, italic, color)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_horizontal_rule(doc, color_hex="B8860B"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr ' + nsdecls("w") + '>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="' + color_hex + '"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def shade_cell(cell, color_hex):
    shading = parse_xml(
        '<w:shd ' + nsdecls("w") + ' w:fill="' + color_hex + '" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def process_inline(text):
    """Strip markdown inline formatting, return clean text."""
    text = re.sub(r'\[\[([^|\]]*?\|)?([^\]]+?)\]\]', r'\2', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    if text.startswith('> '):
        text = text[2:]
    return text


def add_rich_paragraph(doc, text, default_font, default_size,
                       default_color=None, is_italic_block=False):
    """Add paragraph with bold/italic runs parsed from markdown."""
    if default_color is None:
        default_color = BODY_BLACK
    p = doc.add_paragraph()
    text = re.sub(r'\[\[([^|\]]*?\|)?([^\]]+?)\]\]', r'\2', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*]+?\*|_[^_]+?_)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            set_run(run, default_font, default_size, bold=True, italic=True, color=default_color)
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run(run, default_font, default_size, bold=True, color=default_color)
        elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
            run = p.add_run(part[1:-1])
            set_run(run, default_font, default_size, italic=True, color=default_color)
        else:
            run = p.add_run(part)
            set_run(run, default_font, default_size, italic=is_italic_block, color=default_color)
    return p


# ═══════════════════════════════════════════════════════════════
# MARKDOWN PARSING
# ═══════════════════════════════════════════════════════════════
def extract_designation(filename):
    m = re.match(r'(FT-\d+|FP-\d+|SP\d+)', filename)
    return m.group(1) if m else ""


def strip_frontmatter(text):
    if text.startswith('---'):
        end = text.find('---', 3)
        if end != -1:
            fm = text[3:end].strip()
            body = text[end+3:].strip()
            title_match = re.search(r'title:\s*["\']?(.*?)["\']?\s*$', fm, re.MULTILINE)
            title = title_match.group(1) if title_match else None
            return title, body
    return None, text


def extract_title_block(body):
    lines = body.split('\n')
    title = subtitle = author = None
    consumed = 0

    for i, line in enumerate(lines[:15]):
        stripped = line.strip()
        if not stripped or stripped == '---':
            consumed = i + 1
            continue
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:].strip()
            consumed = i + 1
            continue
        if stripped.startswith('### ') and title and not subtitle:
            subtitle = stripped[4:].strip()
            consumed = i + 1
            continue
        if stripped.startswith('_') and stripped.endswith('_') and any(k in stripped for k in ('Lowe', 'Theophysics', '2026', '2025')):
            author = stripped.strip('_').strip()
            consumed = i + 1
            continue
        if ('Lowe' in stripped or 'David Lowe' in stripped) and any(k in stripped for k in ('2026', '2025', 'Theophysics')):
            author = stripped
            consumed = i + 1
            continue
        if title:
            break

    while consumed < len(lines) and (lines[consumed].strip() == '' or lines[consumed].strip() == '---'):
        consumed += 1

    remaining = '\n'.join(lines[consumed:])
    return title, subtitle, author, remaining


# ═══════════════════════════════════════════════════════════════
# CONVERTER
# ═══════════════════════════════════════════════════════════════
def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()

    filename = os.path.basename(md_path)
    designation = extract_designation(filename)

    fm_title, body = strip_frontmatter(raw)
    title, subtitle, author, content = extract_title_block(body)

    if not title and fm_title:
        title = re.sub(r'^DT\d+\s*[\u2014\u2013-]\s*', '', fm_title).strip()
    if not title:
        title = filename.replace('.md', '').split('_', 2)[-1].replace('_', ' ')
    if not author:
        author = "David Lowe | Theophysics Research | March 2026"

    doc = Document()
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # === TITLE BLOCK ===
    add_paragraph(doc, "THEOPHYSICS RESEARCH", "Consolas", Pt(10), bold=True, color=GOLD)
    if designation:
        add_paragraph(doc, designation, "Consolas", Pt(9), color=MED_GRAY, space_after=4)
    add_horizontal_rule(doc)
    add_paragraph(doc, title, "Georgia", Pt(26), bold=True, space_after=4)
    if subtitle:
        add_paragraph(doc, subtitle, "Georgia", Pt(14), italic=True, color=DARK_GRAY, space_after=4)
    author_parts = author.replace(' \u2014 ', ' | ').replace(' - ', ' | ')
    add_paragraph(doc, author_parts, "Georgia", Pt(10), italic=True, color=MED_GRAY, space_after=2)
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # === BODY ===
    lines = content.split('\n')
    i = 0
    table_rows = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                set_run(run, 'Consolas', Pt(9), color=RGBColor(0x33, 0x33, 0x33))
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            if re.match(r'^[\s|:-]+$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            next_is_table = (i + 1 < len(lines) and
                           lines[i+1].strip().startswith('|') and
                           lines[i+1].strip().endswith('|'))
            if not next_is_table and table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < ncols:
                            cell = tbl.cell(ri, ci)
                            cell.text = process_inline(cell_text)
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    set_run(run, 'Arial', Pt(10), color=BODY_BLACK)
                if len(table_rows) > 0:
                    for ci in range(min(ncols, len(tbl.columns))):
                        cell = tbl.cell(0, ci)
                        shade_cell(cell, "F5E6CC")
                        for p in cell.paragraphs:
                            for run in p.runs:
                                set_run(run, 'Arial', Pt(10), bold=True, color=HEADING_COLOR)
                doc.add_paragraph()
                table_rows = []
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ('---', '***', '___'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = process_inline(stripped[level:].strip())
            if level <= 2:
                p = add_paragraph(doc, heading_text, "Georgia", Pt(16), bold=True, color=HEADING_COLOR)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 3:
                p = add_paragraph(doc, heading_text, "Georgia", Pt(13), bold=True, color=ACCENT_BLUE)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            else:
                p = add_paragraph(doc, heading_text, "Arial", Pt(12), bold=True, color=ACCENT_BLUE, italic=True)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            if quote_text.startswith('[!'):
                i += 1
                continue
            p = add_rich_paragraph(doc, quote_text, 'Georgia', Pt(11),
                                   default_color=BLOCKQUOTE_GRAY, is_italic_block=True)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            p = add_rich_paragraph(doc, item_text, 'Georgia', Pt(12), default_color=BODY_BLACK)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            if p.runs:
                p.runs[0].text = "\u2022  " + p.runs[0].text
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            num = m.group(1)
            item_text = m.group(2).strip()
            p = add_rich_paragraph(doc, item_text, 'Georgia', Pt(12), default_color=BODY_BLACK)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            if p.runs:
                p.runs[0].text = f"{num}.  " + p.runs[0].text
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Regular paragraph
        para_lines = [stripped]
        while i + 1 < len(lines):
            next_line = lines[i+1].strip()
            if (not next_line or next_line.startswith('#') or next_line.startswith('- ') or
                next_line.startswith('* ') or next_line.startswith('> ') or
                next_line.startswith('|') or next_line in ('---', '***', '___') or
                next_line.startswith('```') or re.match(r'^\d+\.\s', next_line)):
                break
            para_lines.append(next_line)
            i += 1

        full_text = ' '.join(para_lines)
        p = add_rich_paragraph(doc, full_text, 'Georgia', Pt(12), default_color=BODY_BLACK)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    # === FOOTER ===
    doc.add_paragraph()
    add_horizontal_rule(doc)
    add_paragraph(doc, "Theophysics Research  |  theophysics.pro  |  David Lowe",
                  "Consolas", Pt(8), color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(docx_path)
    return docx_path


# ═══════════════════════════════════════════════════════════════
# PDF EXPORT
# ═══════════════════════════════════════════════════════════════
def export_pdf(docx_path):
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        print("  WARNING: docx2pdf not installed. Run: pip install docx2pdf")
        return None
    except Exception as e:
        print(f"  PDF export failed: {e}")
        return None


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description="Theophysics Paper Engine")
    parser.add_argument('files', nargs='*', help='Specific .md files to convert')
    parser.add_argument('--pdf', action='store_true', help='Also export to PDF')
    parser.add_argument('--dir', default='.', help='Directory to scan for .md files')
    parser.add_argument('--templates', action='store_true', help='Generate 7Q score templates')
    parser.add_argument('--out', default=None, help='Output directory (default: same as input)')
    args = parser.parse_args()

    if args.templates:
        from templates import build_templates
        out_dir = args.out or args.dir
        build_templates(out_dir)
        return

    if args.files:
        md_files = args.files
    else:
        md_files = sorted(
            f for f in os.listdir(args.dir)
            if f.endswith('.md') and (f.startswith('FT-') or f.startswith('FP-') or f.startswith('SP'))
        )

    if not md_files:
        print("No matching .md files found. Pass filenames or use --dir.")
        return

    print(f"Converting {len(md_files)} papers...\n")
    out_dir = args.out or args.dir

    for md_file in md_files:
        md_path = os.path.join(args.dir, md_file) if not os.path.isabs(md_file) else md_file
        base = os.path.basename(md_file).replace('.md', '.docx')
        docx_path = os.path.join(out_dir, base)

        try:
            convert_md_to_docx(md_path, docx_path)
            print(f"  DOCX: {base}")

            if args.pdf:
                pdf = export_pdf(docx_path)
                if pdf:
                    print(f"  PDF:  {os.path.basename(pdf)}")
        except Exception as e:
            print(f"  ERROR: {md_file}: {e}")

    print(f"\nDone. {len(md_files)} papers processed.")


if __name__ == '__main__':
    main()
