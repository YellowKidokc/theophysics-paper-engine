"""
Build two Word document templates for 7Q scored papers:
1. TEMPLATE_A_OPUS.docx - matches Opus/OpenAI scored format
2. TEMPLATE_B_CLAUDE.docx - Claude's own design with colored score boxes
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.dirname(os.path.abspath(__file__))


def build_templates(out_dir=None):
    """Entry point for --templates flag. Builds both templates."""
    global OUT
    if out_dir:
        OUT = out_dir
    build_a()
    build_b()

# === COLORS ===
GOLD = RGBColor(0xB8, 0x86, 0x0B)
MED_GRAY = RGBColor(0x88, 0x88, 0x88)
DARK_GRAY = RGBColor(0x55, 0x55, 0x55)
BODY = RGBColor(0x00, 0x00, 0x00)
NAVY = RGBColor(0x1A, 0x3C, 0x5E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

Q_COLORS = {
    "Q0": {"name": "ARRIVE",       "hex": "8A8D9B", "rgb": RGBColor(0x8A, 0x8D, 0x9B), "bg": "E8E9EC", "icon": "\u25CB"},
    "Q1": {"name": "DEFINE",       "hex": "D4A853", "rgb": RGBColor(0xD4, 0xA8, 0x53), "bg": "FDF5E6", "icon": "\u25C9"},
    "Q2": {"name": "LOCATE",       "hex": "7C6340", "rgb": RGBColor(0x7C, 0x63, 0x40), "bg": "F0EAE0", "icon": "\u25CE"},
    "Q3": {"name": "COMMIT",       "hex": "6B8C42", "rgb": RGBColor(0x6B, 0x8C, 0x42), "bg": "E8F0DD", "icon": "\u25C6"},
    "Q4": {"name": "SUPPORT",      "hex": "38BDF8", "rgb": RGBColor(0x38, 0xBD, 0xF8), "bg": "E0F4FE", "icon": "\u25C7"},
    "Q5": {"name": "GROUND",       "hex": "A0724A", "rgb": RGBColor(0xA0, 0x72, 0x4A), "bg": "F0E6DA", "icon": "\u25BD"},
    "Q6": {"name": "PROPAGATE",    "hex": "34D399", "rgb": RGBColor(0x34, 0xD3, 0x99), "bg": "DCFCE7", "icon": "\u25B3"},
    "Q7": {"name": "FALSIFICATION","hex": "EF4444", "rgb": RGBColor(0xEF, 0x44, 0x44), "bg": "FEE2E2", "icon": "\u2715"},
}
SCORES = {"Q0": 0.80, "Q1": 0.85, "Q2": 0.80, "Q3": 0.70, "Q4": 0.65, "Q5": 0.60, "Q6": 0.75, "Q7": 0.70}


def set_run(run, font, size, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_p(doc, text, font, size, bold=False, italic=False, color=None, align=None, sa=None, sb=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, font, size, bold, italic, color)
    if align: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    if sb is not None: p.paragraph_format.space_before = Pt(sb)
    return p

def gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml('<w:pBdr ' + nsdecls("w") + '><w:bottom w:val="single" w:sz="6" w:space="1" w:color="B8860B"/></w:pBdr>')
    pPr.append(pBdr)

def navy_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml('<w:pBdr ' + nsdecls("w") + '><w:bottom w:val="single" w:sz="8" w:space="1" w:color="1A3C5E"/></w:pBdr>')
    pPr.append(pBdr)

def shade_cell(cell, color_hex):
    shading = parse_xml('<w:shd ' + nsdecls("w") + ' w:fill="' + color_hex + '" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell(cell, text, font, size, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, font, size, bold, color=color)
    if align: p.alignment = align

def set_table_borders(tbl, left_color=None):
    """Set table borders. If left_color given, thick colored left border only."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml("<w:tblPr " + nsdecls("w") + "></w:tblPr>")
        tbl._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    if left_color:
        borders = parse_xml(
            '<w:tblBorders ' + nsdecls("w") + '>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="single" w:sz="18" w:space="0" w:color="' + left_color + '"/>'
            '<w:bottom w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
    else:
        borders = parse_xml(
            '<w:tblBorders ' + nsdecls("w") + '>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
    tblPr.append(borders)


# ═══════════════════════════════════════════════════════════════
# TEMPLATE A: OPUS STYLE
# ═══════════════════════════════════════════════════════════════
def build_a():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    add_p(doc, "THEOPHYSICS RESEARCH", "Consolas", Pt(10), bold=True, color=GOLD)
    add_p(doc, "FP-XXX  |  SCORED PAPER", "Consolas", Pt(9), color=MED_GRAY, sa=2)
    gold_rule(doc)
    add_p(doc, "[Paper Title Here]", "Georgia", Pt(26), bold=True, sa=4)
    add_p(doc, "[Subtitle \u2014 One Line Summary of the Core Claim]", "Georgia", Pt(14), italic=True, color=DARK_GRAY, sa=2)
    add_p(doc, "David Lowe  |  Theophysics Research  |  March 2026", "Georgia", Pt(10), italic=True, color=MED_GRAY)
    gold_rule(doc)

    # Overview strip
    add_p(doc, "OVERVIEW", "Consolas", Pt(9), bold=True, color=GOLD, sb=12, sa=4)
    tbl = doc.add_table(rows=2, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (h, v) in enumerate(zip(
        ["Type", "Confidence", "ISO Status", "T-Score", "T-Enhanced", "Claims"],
        ["Foundational", "Partially Grounded", "ISO-PARALLEL", "0.697", "0.644", "4"]
    )):
        shade_cell(tbl.cell(0, ci), "1A3C5E")
        set_cell(tbl.cell(0, ci), h, "Arial", Pt(8), bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(tbl.cell(1, ci), "F5F0E8")
        set_cell(tbl.cell(1, ci), v, "Arial", Pt(9), color=BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Score bar
    add_p(doc, "7Q SCORE PROFILE", "Consolas", Pt(9), bold=True, color=GOLD, sa=4)
    bar = doc.add_table(rows=2, cols=8)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (qk, qv) in enumerate(Q_COLORS.items()):
        shade_cell(bar.cell(0, ci), qv["hex"])
        set_cell(bar.cell(0, ci), qk, "Consolas", Pt(8), bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(bar.cell(1, ci), qv["bg"])
        set_cell(bar.cell(1, ci), f"{SCORES[qk]:.2f}", "Consolas", Pt(10), bold=True, color=qv["rgb"], align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    gold_rule(doc)

    # Paper body placeholder
    add_p(doc, "[PAPER CONTENT]", "Georgia", Pt(16), bold=True, color=NAVY, sb=14, sa=8)
    add_p(doc, "[Full paper text goes here \u2014 untouched by the scoring engine. The paper stands on its own above; the 7Q analysis follows below.]", "Georgia", Pt(12), italic=True, color=DARK_GRAY, sa=12)
    gold_rule(doc)

    # 7Q Detailed Analysis
    add_p(doc, "7Q DETAILED ANALYSIS", "Consolas", Pt(11), bold=True, color=GOLD, sb=14, sa=8)

    for qk, qv in Q_COLORS.items():
        score = SCORES[qk]
        hdr = doc.add_table(rows=1, cols=1)
        hdr.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = hdr.cell(0, 0)
        shade_cell(cell, qv["hex"])
        set_cell(cell, f"  {qv['icon']}  {qk} \u2014 {qv['name']}  ({score:.2f})", "Consolas", Pt(11), bold=True, color=WHITE)
        cell.width = Inches(6.5)
        set_table_borders(hdr)
        add_p(doc, f"[{qk} analysis content \u2014 assessment, evidence, caveats]", "Georgia", Pt(11), italic=True, color=DARK_GRAY, sa=4)
        doc.add_paragraph()

    gold_rule(doc)
    add_p(doc, "EXECUTIVE SUMMARY", "Consolas", Pt(11), bold=True, color=GOLD, sb=8, sa=6)
    add_p(doc, "[One-paragraph verdict summarizing what survives scrutiny, what doesn't, and what remains untested.]", "Georgia", Pt(12), italic=True, color=DARK_GRAY, sa=8)

    # Theory map
    add_p(doc, "THEORY RESONANCE MAP", "Consolas", Pt(9), bold=True, color=RGBColor(0xA7, 0x8B, 0xFA), sa=4)
    tm = doc.add_table(rows=4, cols=4)
    tm.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Theory", "Mapping", "Status", "Upgrade Path"]):
        shade_cell(tm.cell(0, ci), "1A3C5E")
        set_cell(tm.cell(0, ci), h, "Arial", Pt(8), bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri in range(1, 4):
        for ci in range(4):
            shade_cell(tm.cell(ri, ci), "F5F0E8")
            set_cell(tm.cell(ri, ci), "[...]", "Arial", Pt(9), color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    gold_rule(doc)
    add_p(doc, "Scored by 7Q Engine v1.0  |  Theophysics Research  |  March 2026", "Consolas", Pt(8), color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(OUT, "TEMPLATE_A_OPUS.docx"))
    print("TEMPLATE_A_OPUS.docx created")


# ═══════════════════════════════════════════════════════════════
# TEMPLATE B: CLAUDE DESIGN - academic, colored left-border boxes
# ═══════════════════════════════════════════════════════════════
def build_b():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Thin gold accent bar
    bar = doc.add_table(rows=1, cols=1)
    cell = bar.cell(0, 0)
    shade_cell(cell, "B8860B")
    set_cell(cell, " ", "Arial", Pt(2))
    set_table_borders(bar)

    add_p(doc, "THEOPHYSICS RESEARCH PAPERS", "Consolas", Pt(8), color=GOLD, sb=8, sa=0)
    add_p(doc, "[Paper Title Here]", "Georgia", Pt(24), bold=True, sa=2)
    add_p(doc, "[Subtitle]", "Georgia", Pt(13), italic=True, color=DARK_GRAY, sa=2)
    add_p(doc, "David Lowe", "Georgia", Pt(11), color=BODY, sa=0)
    add_p(doc, "Theophysics Research  |  theophysics.pro  |  March 2026", "Georgia", Pt(9), color=MED_GRAY, sa=4)
    navy_rule(doc)

    # Score card: Q scores left, metadata right
    add_p(doc, "SCORE CARD", "Consolas", Pt(9), bold=True, color=NAVY, sb=10, sa=6)
    sc = doc.add_table(rows=10, cols=5)
    sc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Dimension", "Score", "  ", "Metric", "Value"]):
        shade_cell(sc.cell(0, ci), "1A3C5E")
        set_cell(sc.cell(0, ci), h, "Consolas", Pt(8), bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, (qk, qv) in enumerate(Q_COLORS.items()):
        row = ri + 1
        shade_cell(sc.cell(row, 0), qv["bg"])
        set_cell(sc.cell(row, 0), f"{qv['icon']} {qk} {qv['name']}", "Consolas", Pt(8), bold=True, color=qv["rgb"])
        shade_cell(sc.cell(row, 1), qv["bg"])
        set_cell(sc.cell(row, 1), f"{SCORES[qk]:.2f}", "Consolas", Pt(10), bold=True, color=qv["rgb"], align=WD_ALIGN_PARAGRAPH.CENTER)

    meta = [("Type", "Foundational"), ("Confidence", "Partially Grounded"), ("ISO Status", "ISO-PARALLEL"),
            ("T-Score", "0.697"), ("T-Enhanced", "0.644"), ("Claims", "4"), ("Strongest", "Q1"), ("Weakest", "Q5")]
    for ri, (label, val) in enumerate(meta):
        row = ri + 1
        if row < 10:
            set_cell(sc.cell(row, 3), label, "Consolas", Pt(8), color=MED_GRAY)
            set_cell(sc.cell(row, 4), val, "Consolas", Pt(9), bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    shade_cell(sc.cell(9, 0), "1A3C5E")
    set_cell(sc.cell(9, 0), "COMPOSITE", "Consolas", Pt(8), bold=True, color=WHITE)
    shade_cell(sc.cell(9, 1), "1A3C5E")
    set_cell(sc.cell(9, 1), "0.697", "Consolas", Pt(11), bold=True, color=RGBColor(0xD4, 0xA8, 0x53), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    navy_rule(doc)

    # Paper body
    add_p(doc, "[PAPER CONTENT]", "Georgia", Pt(16), bold=True, color=NAVY, sb=10, sa=6)
    add_p(doc, "[Full paper text \u2014 stands alone, unmodified by the scoring engine]", "Georgia", Pt(11), italic=True, color=DARK_GRAY, sa=10)
    navy_rule(doc)

    # 7Q Analysis with colored left-border boxes
    add_p(doc, "7Q DIMENSIONAL ANALYSIS", "Consolas", Pt(10), bold=True, color=NAVY, sb=10, sa=8)

    for qk, qv in Q_COLORS.items():
        score = SCORES[qk]
        box = doc.add_table(rows=2, cols=1)
        box.alignment = WD_TABLE_ALIGNMENT.LEFT

        title_cell = box.cell(0, 0)
        shade_cell(title_cell, qv["bg"])
        title_cell.width = Inches(6.5)
        p = title_cell.paragraphs[0]
        r1 = p.add_run(f"  {qv['icon']}  {qk} \u2014 {qv['name']}  ")
        set_run(r1, "Consolas", Pt(10), bold=True, color=qv["rgb"])
        r2 = p.add_run(f"  {score:.2f}")
        set_run(r2, "Consolas", Pt(12), bold=True, color=qv["rgb"])

        content_cell = box.cell(1, 0)
        content_cell.width = Inches(6.5)
        cp = content_cell.paragraphs[0]
        cr = cp.add_run(f"[{qk} assessment: posture evaluation, evidence cited, caveats noted, competing explanations addressed]")
        set_run(cr, "Georgia", Pt(10), italic=True, color=DARK_GRAY)

        set_table_borders(box, left_color=qv["hex"])
        add_p(doc, "", "Arial", Pt(4), sa=2)

    # Verdict
    navy_rule(doc)
    add_p(doc, "VERDICT", "Consolas", Pt(10), bold=True, color=GOLD, sb=8, sa=6)
    add_p(doc, "[Executive summary \u2014 what survives, what dies, what remains untested]", "Georgia", Pt(12), italic=True, color=DARK_GRAY, sa=10)

    # Footer
    gold_rule(doc)
    add_p(doc, "7Q Engine v1.0  |  Theophysics Research  |  theophysics.pro", "Consolas", Pt(8), color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(os.path.join(OUT, "TEMPLATE_B_CLAUDE.docx"))
    print("TEMPLATE_B_CLAUDE.docx created")


build_a()
build_b()
print("\nBoth templates saved to:", OUT)
